import json
import os
from datetime import datetime
from io import BytesIO

import streamlit as st
from langchain_core.prompts import ChatPromptTemplate
from langchain_groq import ChatGroq

# ============ IMPORTS & SETUP ============
try:
    import PyPDF2
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    from docx import Document as DocxDocument
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    PDF_EXPORT_SUPPORT = True
except ImportError:
    PDF_EXPORT_SUPPORT = False

# ============ PAGE CONFIG ============
st.set_page_config(
    page_title="Resumake",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# ============ LLM INITIALIZATION ============
@st.cache_resource
def init_llm(api_key):
    if not api_key:
        return None
    return ChatGroq(
        temperature=0.3,
        model_name="llama-3.3-70b-versatile",
        groq_api_key=api_key
    )

def get_api_key():
    api_key = os.getenv("GROQ_API_KEY")
    if api_key:
        return api_key
    try:
        return st.secrets["GROQ_API_KEY"]
    except (FileNotFoundError, KeyError):
        pass
    return ""

initial_api_key = get_api_key()
llm = init_llm(initial_api_key) if initial_api_key else None

# ============ SESSION STATE ============
if 'resume_data' not in st.session_state:
    st.session_state.resume_data = {
        'personal_info': {},
        'education': [],
        'certifications': [],
        'experience': [],
        'projects': [],
        'skills': []
    }
if 'generated_resume' not in st.session_state:
    st.session_state.generated_resume = ""
if 'resume_versions' not in st.session_state:
    st.session_state.resume_versions = {}
if 'ats_score' not in st.session_state:
    st.session_state.ats_score = None
if 'ats_feedback' not in st.session_state:
    st.session_state.ats_feedback = ""
if 'edit_mode' not in st.session_state:
    st.session_state.edit_mode = False
if 'page' not in st.session_state:
    st.session_state.page = "Home"
if 'chat_history_general' not in st.session_state:
    st.session_state.chat_history_general = []
if 'chat_history_guidance' not in st.session_state:
    st.session_state.chat_history_guidance = []
if 'chat_history_interview' not in st.session_state:
    st.session_state.chat_history_interview = []
if 'uploaded_resume' not in st.session_state:
    st.session_state.uploaded_resume = ""

# ============ EXPORT FUNCTIONS ============
def export_to_pdf(resume_text):
    """Export resume to PDF format"""
    try:
        if not PDF_EXPORT_SUPPORT:
            return None
        
        from reportlab.lib.enums import TA_LEFT
        
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch, leftMargin=0.5*inch, rightMargin=0.5*inch)
        styles = getSampleStyleSheet()
        story = []
        
        for line in resume_text.split('\n'):
            line = line.strip()
            if not line:
                story.append(Spacer(1, 0.05*inch))
                continue
            
            line_safe = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            
            if line.isupper() and len(line.split()) <= 8:
                para = Paragraph(f"<b>{line_safe}</b>", styles['Normal'])
            else:
                para = Paragraph(line_safe, styles['Normal'])
            
            story.append(para)
        
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    except Exception as e:
        st.error(f"PDF Error: {str(e)}")
        return None

def export_to_docx(resume_text):
    """Export resume to DOCX format"""
    try:
        if not DOCX_SUPPORT:
            return None
        
        from docx import Document
        from docx.shared import Pt
        
        doc = Document()
        
        for line in resume_text.split('\n'):
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue
            
            p = doc.add_paragraph(line)
            
            if line.isupper() and len(line.split()) <= 8:
                for run in p.runs:
                    run.font.bold = True
                    run.font.size = Pt(12)
            else:
                for run in p.runs:
                    run.font.size = Pt(11)
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    except Exception as e:
        st.error(f"DOCX Error: {str(e)}")
        return None

# ============ RESUME TEMPLATES ============
RESUME_TEMPLATES = {
    "Professional": "Create a formal, professional resume suitable for corporate positions.",
    "Creative": "Create a modern, creative resume with emphasis on design and innovation skills.",
    "Technical": "Create a technical resume highlighting programming languages, tools, and technical projects.",
    "Academic": "Create an academic resume emphasizing research, publications, and education.",
    "Hybrid": "Create a hybrid resume combining both chronological and functional formats."
}

# ============ SIDEBAR ============
with st.sidebar:
    st.title("🤖 RESUMAKE")
    
    has_api_key = bool(initial_api_key)
    
    if not has_api_key:
        st.warning("⚠️ API Key Required")
        with st.expander("Configure API Key"):
            groq_api_key = st.text_input("Enter Groq API Key", type="password", help="Get one at https://console.groq.com")
            if groq_api_key:
                os.environ["GROQ_API_KEY"] = groq_api_key
                llm = init_llm(groq_api_key)
                st.success("✅ API Key saved!")
                st.rerun()
    
    st.markdown("---")
    st.markdown("### 📂 Select Option")
    
    pages = ["Home", "Resume Builder", "ATS Scanner", "AI Assistant"]
    for page_name in pages:
        emoji = {'Home': '🏠', 'Resume Builder': '📝', 'ATS Scanner': '🔍', 'AI Assistant': '💬'}.get(page_name, '📄')
        if st.button(f"{emoji} {page_name}", use_container_width=True, type="primary" if st.session_state.page == page_name else "secondary"):
            st.session_state.page = page_name
            st.rerun()
    
    st.markdown("---")
    st.markdown(f"- TXT: ✅\n- PDF: {'✅' if PDF_EXPORT_SUPPORT else '❌'}\n- DOCX: {'✅' if DOCX_SUPPORT else '❌'}")
    st.caption("Made with ❤️ using Groq Cloud + LangChain")

# ============ MAIN PAGES ============

if st.session_state.page == "Home":
    st.markdown("<h1 style='text-align: center; color: #1f77b4; font-size: 3.5em;'>🤖 RESUMAKE</h1>", unsafe_allow_html=True)
    st.markdown("<p style='text-align: center; font-size: 1.3em; color: #999;'>Build professional resumes and optimize them for Applicant Tracking Systems</p>", unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    with col1:
        st.markdown("""
        <div style='padding: 30px; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); border-radius: 15px; color: white;'>
            <h2 style='color: white; margin-top: 0;'>📝 Resume Builder</h2>
            <ul style='font-size: 1.05em; line-height: 1.8;'>
                <li>✨ Multiple resume templates</li>
                <li>🤖 AI-powered generation</li>
                <li>🔄 Version control</li>
                <li>📥 Multi-format download</li>
            </ul>
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        st.markdown("""
        <div style='padding: 30px; background: linear-gradient(135deg, #f093fb 0%, #f5576c 100%); border-radius: 15px; color: white;'>
            <h2 style='color: white; margin-top: 0;'>🔍 ATS Scanner</h2>
            <ul style='font-size: 1.05em; line-height: 1.8;'>
                <li>📊 Detailed scoring</li>
                <li>🎯 Keyword matching</li>
                <li>💡 Smart suggestions</li>
                <li>📈 Improvement tracking</li>
            </ul>
        </div>
        """, unsafe_allow_html=True)
    st.markdown("<h2 style='text-align: center; color: #fff; font-size: 2.5em;'>🚀 How It Works</h2>", unsafe_allow_html=True)
    st.markdown("<br>", unsafe_allow_html=True)
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.markdown("""
        <div style='text-align: center; padding: 25px; background: #f8f9fa; border-radius: 12px; height: 100%;'>
            <div style='font-size: 3em; margin-bottom: 15px;'>1️⃣</div>
            <h3 style='color: #1f77b4; margin-bottom: 15px;'>Choose Your Tool</h3>
            <p style='color: #666; font-size: 1.05em;'>Select Resume Builder to create a new resume or ATS Scanner to analyze an existing one from the sidebar</p>
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        st.markdown("""
        <div style='text-align: center; padding: 25px; background: #f8f9fa; border-radius: 12px; height: 100%;'>
            <div style='font-size: 3em; margin-bottom: 15px;'>2️⃣</div>
            <h3 style='color: #1f77b4; margin-bottom: 15px;'>Input Your Data</h3>
            <p style='color: #666; font-size: 1.05em;'>Fill in your information or upload your resume. Our AI will process and optimize it for maximum impact</p>
        </div>
        """, unsafe_allow_html=True)
    
    with col3:
        st.markdown("""
        <div style='text-align: center; padding: 25px; background: #f8f9fa; border-radius: 12px; height: 100%;'>
            <div style='font-size: 3em; margin-bottom: 15px;'>3️⃣</div>
            <h3 style='color: #1f77b4; margin-bottom: 15px;'>Get Results</h3>
            <p style='color: #666; font-size: 1.05em;'>Download your polished resume or receive detailed ATS analysis with actionable improvements</p>
        </div>
        """, unsafe_allow_html=True)
    
    st.markdown("<br><br>", unsafe_allow_html=True)
    
    # Key features section
    st.markdown("<h2 style='text-align: center; color: #fff; font-size: 2.5em;'>✨ Key Features</h2>", unsafe_allow_html=True)
    st.markdown("<br>", unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("""
        <div style='padding: 20px; background: white; border-left: 4px solid #667eea; margin-bottom: 15px; border-radius: 8px; box-shadow: 0 2px 4px rgba(0,0,0,0.05);'>
            <h4 style='color: #667eea; margin-top: 0;'>🎯 AI-Powered Generation</h4>
            <p style='color: #666;'>Leverage advanced AI to create compelling resume content with strong action verbs and quantifiable achievements</p>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("""
        <div style='padding: 20px; background: white; border-left: 4px solid #667eea; margin-bottom: 15px; border-radius: 8px; box-shadow: 0 2px 4px rgba(0,0,0,0.05);'>
            <h4 style='color: #667eea; margin-top: 0;'>📊 ATS Compatibility</h4>
            <p style='color: #666;'>Ensure your resume passes Applicant Tracking Systems with optimized formatting and keywords</p>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("""
        <div style='padding: 20px; background: white; border-left: 4px solid #667eea; margin-bottom: 15px; border-radius: 8px; box-shadow: 0 2px 4px rgba(0,0,0,0.05);'>
            <h4 style='color: #667eea; margin-top: 0;'>✏️ Easy Editing</h4>
            <p style='color: #666;'>Edit and customize your AI-generated resume to perfectly match your style and preferences</p>
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        st.markdown("""
        <div style='padding: 20px; background: white; border-left: 4px solid #f5576c; margin-bottom: 15px; border-radius: 8px; box-shadow: 0 2px 4px rgba(0,0,0,0.05);'>
            <h4 style='color: #f5576c; margin-top: 0;'>📈 Score Analysis</h4>
            <p style='color: #666;'>Get a detailed ATS compatibility score with specific feedback on how to improve your resume</p>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("""
        <div style='padding: 20px; background: white; border-left: 4px solid #f5576c; margin-bottom: 15px; border-radius: 8px; box-shadow: 0 2px 4px rgba(0,0,0,0.05);'>
            <h4 style='color: #f5576c; margin-top: 0;'>📄 Multiple Formats</h4>
            <p style='color: #666;'>Upload and download resumes in various formats including PDF, DOCX, and TXT</p>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("""
        <div style='padding: 20px; background: white; border-left: 4px solid #f5576c; margin-bottom: 15px; border-radius: 8px; box-shadow: 0 2px 4px rgba(0,0,0,0.05);'>
            <h4 style='color: #f5576c; margin-top: 0;'>🔍 Keyword Matching</h4>
            <p style='color: #666;'>Identify missing keywords from job descriptions and get suggestions to improve your match rate</p>
        </div>
        """, unsafe_allow_html=True)
    
    st.markdown("<br><br>", unsafe_allow_html=True)

    # Call to action
    st.markdown("""
    <div style='text-align: center; padding: 40px; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); border-radius: 15px; color: white;'>
        <h2 style='color: white; font-size: 2.2em; margin-top: 0;'>Ready to Get Started?</h2>
        <p style='font-size: 1.2em; margin-bottom: 25px;'>Choose an option from the sidebar to begin building or scanning your resume!</p>
        <p style='font-size: 1.1em; opacity: 0.9;'>👈 Select <strong>Resume Builder</strong> to create a new resume<br>or <strong>ATS Scanner</strong> to analyze an existing one</p>
    </div>
    """, unsafe_allow_html=True)

elif st.session_state.page == "Resume Builder":
    st.title("📝 Resume Builder")
    
    st.markdown("### 🎨 Choose Resume Template")
    selected_template = st.selectbox(
        "Select a template style:",
        list(RESUME_TEMPLATES.keys()),
        help="Choose a template that best fits your career type"
    )
    st.info(f"📋 {RESUME_TEMPLATES[selected_template]}")
    
    with st.expander("👤 Personal Information", expanded=True):
        col1, col2 = st.columns(2)
        with col1:
            name = st.text_input("Full Name*", key="name_input")
            email = st.text_input("Email*", key="email_input")
            phone = st.text_input("Phone", key="phone_input")
        with col2:
            linkedin = st.text_input("LinkedIn URL", key="linkedin_input")
            github = st.text_input("GitHub URL", key="github_input")
            portfolio = st.text_input("Portfolio URL", key="portfolio_input")
        
        target_role = st.text_input("Target Role*", key="target_input")
        summary = st.text_area("Professional Summary", key="summary_input")
        
        if st.button("Save Personal Info", key="save_personal"):
            st.session_state.resume_data['personal_info'] = {
                'name': name, 'email': email, 'phone': phone,
                'linkedin': linkedin, 'github': github, 'portfolio': portfolio,
                'target_role': target_role, 'summary': summary
            }
            st.success("✅ Personal information saved!")
    
    with st.expander("🎓 Education"):
        num_edu = st.number_input("Number of Education Entries", min_value=0, max_value=10, value=1, key="num_edu")
        education_entries = []
        
        for i in range(num_edu):
            st.markdown(f"**Education #{i+1}**")
            col1, col2 = st.columns(2)
            with col1:
                degree = st.text_input(f"Degree", key=f"edu_degree_{i}")
                institution = st.text_input(f"Institution", key=f"edu_inst_{i}")
            with col2:
                year = st.text_input(f"Year", key=f"edu_year_{i}")
                gpa = st.text_input(f"GPA (optional)", key=f"edu_gpa_{i}")
            
            if degree and institution:
                education_entries.append({'degree': degree, 'institution': institution, 'year': year, 'gpa': gpa})
            st.markdown("---")
        
        if st.button("Save Education", key="save_edu"):
            st.session_state.resume_data['education'] = education_entries
            st.success("✅ Education saved!")
    
    with st.expander("🏆 Certifications"):
        num_certs = st.number_input("Number of Certifications", min_value=0, max_value=10, value=1, key="num_certs")
        cert_entries = []
        
        for i in range(num_certs):
            st.markdown(f"**Certification #{i+1}**")
            col1, col2, col3 = st.columns(3)
            with col1:
                cert_name = st.text_input(f"Certification Name", key=f"cert_name_{i}")
            with col2:
                cert_org = st.text_input(f"Issuing Organization", key=f"cert_org_{i}")
            with col3:
                cert_date = st.text_input(f"Date", key=f"cert_date_{i}")
            
            if cert_name and cert_org:
                cert_entries.append({'name': cert_name, 'organization': cert_org, 'date': cert_date})
            st.markdown("---")
        
        if st.button("Save Certifications", key="save_certs"):
            st.session_state.resume_data['certifications'] = cert_entries
            st.success("✅ Certifications saved!")
    
    with st.expander("💼 Work Experience"):
        num_exp = st.number_input("Number of Work Experiences", min_value=0, max_value=10, value=1, key="num_exp")
        exp_entries = []
        
        for i in range(num_exp):
            st.markdown(f"**Experience #{i+1}**")
            col1, col2 = st.columns(2)
            with col1:
                title = st.text_input(f"Job Title", key=f"exp_title_{i}")
                company = st.text_input(f"Company", key=f"exp_company_{i}")
            with col2:
                duration = st.text_input(f"Duration", key=f"exp_duration_{i}")
            
            description = st.text_area(f"Key Achievements", key=f"exp_desc_{i}", height=120)
            
            if title and company:
                exp_entries.append({'title': title, 'company': company, 'duration': duration, 'description': description})
            st.markdown("---")
        
        if st.button("Save Work Experience", key="save_exp"):
            st.session_state.resume_data['experience'] = exp_entries
            st.success("✅ Work experience saved!")
    
    with st.expander("🚀 Projects"):
        num_projects = st.number_input("Number of Projects", min_value=0, max_value=10, value=1, key="num_projects")
        project_entries = []
        
        for i in range(num_projects):
            st.markdown(f"**Project #{i+1}**")
            project_title = st.text_input(f"Project Title", key=f"proj_title_{i}")
            project_link = st.text_input(f"Project Link (optional)", key=f"proj_link_{i}")
            technologies = st.text_input(f"Technologies Used", key=f"proj_tech_{i}")
            achievements = st.text_area(f"Key Achievements", key=f"proj_achievements_{i}", height=100)
            
            if project_title and technologies:
                project_entries.append({
                    'title': project_title,
                    'link': project_link,
                    'technologies': technologies,
                    'achievements': achievements
                })
            st.markdown("---")
        
        if st.button("Save Projects", key="save_projects"):
            st.session_state.resume_data['projects'] = project_entries
            st.success("✅ Projects saved!")
    
    with st.expander("🛠️ Skills"):
        num_skills = st.number_input("Number of Skill Categories", min_value=0, max_value=10, value=1, key="num_skills")
        skill_entries = []
        
        for i in range(num_skills):
            st.markdown(f"**Skill Category #{i+1}**")
            col1, col2 = st.columns(2)
            with col1:
                category = st.text_input(f"Category", key=f"skill_cat_{i}")
            with col2:
                items = st.text_input(f"Skills (comma separated)", key=f"skill_items_{i}")
            
            if category and items:
                skill_entries.append({'category': category, 'items': items})
            st.markdown("---")
        
        if st.button("Save Skills", key="save_skills"):
            st.session_state.resume_data['skills'] = skill_entries
            st.success("✅ Skills saved!")
    
    st.markdown("---")
    
    col1, col2, col3 = st.columns([1, 1, 1])
    
    with col1:
        if st.button("🚀 Generate Resume", type="primary", use_container_width=True, key="generate_btn"):
            if not initial_api_key or not llm:
                st.error("⚠️ Please configure your Groq API key!")
            elif not st.session_state.resume_data['personal_info'].get('name'):
                st.error("Please fill in personal information!")
            else:
                with st.spinner("Generating your resume..."):
                    resume_prompt = ChatPromptTemplate.from_messages([
                        ("system", f"""You are an expert resume writer. Create a {selected_template} style resume.
                        
Template guidelines: {RESUME_TEMPLATES[selected_template]}

Format with clear sections, use strong action verbs, include metrics, and ensure ATS compatibility."""),
                        ("human", """Create a professional resume with this information:

Personal: {personal_info}
Education: {education}
Certifications: {certifications}
Experience: {experience}
Projects: {projects}
Skills: {skills}
Target Role: {target_role}""")
                    ])
                    
                    resume_data = st.session_state.resume_data
                    chain = resume_prompt | llm
                    response = chain.invoke({
                        "personal_info": json.dumps(resume_data['personal_info'], indent=2),
                        "education": json.dumps(resume_data['education'], indent=2),
                        "certifications": json.dumps(resume_data['certifications'], indent=2),
                        "experience": json.dumps(resume_data['experience'], indent=2),
                        "projects": json.dumps(resume_data['projects'], indent=2),
                        "skills": json.dumps(resume_data['skills'], indent=2),
                        "target_role": resume_data['personal_info'].get('target_role', 'Professional')
                    })
                    
                    st.session_state.generated_resume = response.content
                    st.session_state.edit_mode = False
                    st.success("✅ Resume generated!")
                    st.rerun()
    
    with col2:
        if st.button("📋 AI Review", use_container_width=True, key="review_btn"):
            if not st.session_state.generated_resume:
                st.error("Please generate a resume first!")
            else:
                with st.spinner("Reviewing..."):
                    review_prompt = ChatPromptTemplate.from_messages([
                        ("system", "You are an expert resume reviewer. Provide brief, actionable feedback."),
                        ("human", "Review this resume and suggest 3-5 key improvements:\n{resume}")
                    ])
                    
                    chain = review_prompt | llm
                    response = chain.invoke({"resume": st.session_state.generated_resume})
                    
                    st.info(response.content)
    
    with col3:
        if st.button("💾 Save Version", use_container_width=True, key="save_version_btn"):
            if not st.session_state.generated_resume:
                st.error("Please generate a resume first!")
            else:
                timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")
                version_name = f"Version_{len(st.session_state.resume_versions) + 1}_{timestamp}"
                st.session_state.resume_versions[version_name] = st.session_state.generated_resume
                st.success(f"✅ Saved as {version_name}")
    
    if st.session_state.generated_resume:
        st.markdown("---")
        st.subheader("📄 Generated Resume")
        
        col1, col2 = st.columns([1, 5])
        with col1:
            if st.button("✏️ Edit" if not st.session_state.edit_mode else "👁️ View"):
                st.session_state.edit_mode = not st.session_state.edit_mode
                st.rerun()
        
        if st.session_state.edit_mode:
            edited_resume = st.text_area("Edit Your Resume", value=st.session_state.generated_resume, height=600, key="edit_resume")
            if st.button("💾 Save Changes"):
                st.session_state.generated_resume = edited_resume
                st.session_state.edit_mode = False
                st.success("✅ Changes saved!")
                st.rerun()
        else:
            st.text_area("Resume Preview", value=st.session_state.generated_resume, height=600, disabled=True, key="view_resume")
        
        st.markdown("### 📊 Resume Statistics")
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            word_count = len(st.session_state.generated_resume.split())
            st.metric("Word Count", word_count)
        
        with col2:
            line_count = len(st.session_state.generated_resume.split('\n'))
            st.metric("Lines", line_count)
        
        with col3:
            action_verbs = sum(1 for verb in ['led', 'managed', 'developed', 'implemented', 'created', 'improved'] 
                              if verb in st.session_state.generated_resume.lower())
            st.metric("Action Verbs", action_verbs)
        
        with col4:
            metrics_count = sum(1 for char in st.session_state.generated_resume if char.isdigit())
            st.metric("Metrics/Numbers", metrics_count)
        
        st.markdown("### 💾 Download Resume")
        file_name_base = st.session_state.resume_data['personal_info'].get('name', 'resume').replace(' ', '_')
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.download_button(
                label="📥 Download as TXT",
                data=st.session_state.generated_resume,
                file_name=f"{file_name_base}_resume.txt",
                mime="text/plain",
                use_container_width=True,
                key="download_txt"
            )
        
        with col2:
            if PDF_EXPORT_SUPPORT:
                pdf_data = export_to_pdf(st.session_state.generated_resume)
                if pdf_data:
                    st.download_button(
                        label="📥 Download as PDF",
                        data=pdf_data,
                        file_name=f"{file_name_base}_resume.pdf",
                        mime="application/pdf",
                        use_container_width=True,
                        key="download_pdf"
                    )
        
        with col3:
            if DOCX_SUPPORT:
                docx_data = export_to_docx(st.session_state.generated_resume)
                if docx_data:
                    st.download_button(
                        label="📥 Download as DOCX",
                        data=docx_data,
                        file_name=f"{file_name_base}_resume.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                        key="download_docx"
                    )
        
        if st.session_state.resume_versions:
            st.markdown("### 📚 Version History")
            version_to_load = st.selectbox("Load a previous version:", list(st.session_state.resume_versions.keys()))
            if st.button("Load Selected Version"):
                st.session_state.generated_resume = st.session_state.resume_versions[version_to_load]
                st.success(f"Loaded {version_to_load}")
                st.rerun()

elif st.session_state.page == "ATS Scanner":
    st.title("🔍 ATS Resume Scanner")
    
    if not initial_api_key or not llm:
        st.error("⚠️ Please configure your Groq API key!")
    else:
        resume_source = st.radio("Choose Resume Source:", ["Use Generated Resume", "Upload File", "Paste Text"], horizontal=True)
        
        resume_to_scan = ""
        
        if resume_source == "Use Generated Resume":
            if not st.session_state.generated_resume:
                st.warning("⚠️ No generated resume found.")
            else:
                resume_to_scan = st.session_state.generated_resume
                st.success("✅ Using generated resume")
        
        elif resume_source == "Upload File":
            uploaded_file = st.file_uploader("Upload resume", type=['txt', 'pdf', 'docx'])
            if uploaded_file:
                try:
                    if uploaded_file.type == "text/plain":
                        resume_to_scan = uploaded_file.read().decode("utf-8")
                    elif uploaded_file.type == "application/pdf" and PDF_SUPPORT:
                        pdf_reader = PyPDF2.PdfReader(uploaded_file)
                        resume_to_scan = "\n".join([page.extract_text() for page in pdf_reader.pages])
                    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" and DOCX_SUPPORT:
                        doc = DocxDocument(uploaded_file)
                        resume_to_scan = "\n".join([para.text for para in doc.paragraphs])
                    st.success(f"✅ File uploaded: {uploaded_file.name}")
                except Exception as e:
                    st.error(f"Error reading file: {str(e)}")
        
        else:
            pasted_resume = st.text_area("Paste Your Resume", height=300, key="paste_resume")
            if pasted_resume:
                resume_to_scan = pasted_resume
        
        st.markdown("---")
        
        st.markdown("### 📋 Job Descriptions")
        num_jobs = st.number_input("Number of job descriptions to analyze:", min_value=1, max_value=3, value=1, key="num_jobs")
        
        job_descriptions = []
        for i in range(num_jobs):
            st.markdown(f"**Job Description #{i+1}**")
            job_desc = st.text_area(f"Paste job description {i+1}:", height=200, key=f"job_desc_{i}")
            if job_desc:
                job_descriptions.append(job_desc)
        
        compare_mode = st.checkbox("Compare across all job descriptions", value=False)
        
        if st.button("🔍 Scan Resume", type="primary", use_container_width=True):
            if not resume_to_scan:
                st.error("Please provide a resume!")
            elif not job_descriptions:
                st.error("Please provide at least one job description!")
            else:
                for idx, job_desc in enumerate(job_descriptions):
                    st.markdown(f"### 📊 Analysis for Job #{idx+1}")
                    
                    with st.spinner(f"Analyzing job description {idx+1}..."):
                        ats_prompt = ChatPromptTemplate.from_messages([
                            ("system", """You are an ATS analyzer. Provide detailed analysis with:
1. ATS Score (0-100)
2. Keyword matching percentage
3. Missing keywords
4. Format issues
5. Specific improvement suggestions

Respond in JSON with 'score', 'keyword_match', 'missing_keywords', 'format_issues', and 'suggestions' fields."""),
                            ("human", """Job Description:\n{job_description}\n\nResume:\n{resume}\n\nProvide comprehensive analysis.""")
                        ])
                        
                        chain = ats_prompt | llm
                        response = chain.invoke({"job_description": job_desc, "resume": resume_to_scan})
                        
                        try:
                            response_text = response.content
                            json_start = response_text.find('{')
                            json_end = response_text.rfind('}') + 1
                            result = json.loads(response_text[json_start:json_end])
                            
                            score = result.get('score', 75)
                            keyword_match = result.get('keyword_match', '60%')
                            missing_keywords = result.get('missing_keywords', [])
                            suggestions = result.get('suggestions', response_text)
                            
                            color = "#00CC44" if score >= 80 else "#FFA500" if score >= 60 else "#FFD700" if score >= 40 else "#FF4444"
                            status = ["Needs Improvement ❌", "Fair ⚠️", "Good 👍", "Excellent ✅"][min(3, score // 40)]
                            
                            progress_html = f"""
                            <div style='position: relative; width: 100%; height: 50px; background-color: #f0f0f0; border-radius: 10px; overflow: hidden; margin: 20px 0;'>
                                <div style='width: {score}%; height: 100%; background: linear-gradient(90deg, {color} 0%, {color} 100%);'></div>
                                <div style='position: absolute; top: 50%; left: 50%; transform: translate(-50%, -50%); font-weight: bold; font-size: 18px; color: #333;'>
                                    Score: {score} — {status}
                                </div>
                            </div>
                            """
                            st.markdown(progress_html, unsafe_allow_html=True)
                            
                            col1, col2, col3 = st.columns(3)
                            with col1:
                                st.metric("ATS Score", score)
                            with col2:
                                st.metric("Keyword Match", keyword_match)
                            with col3:
                                st.metric("Missing Keywords", len(missing_keywords) if isinstance(missing_keywords, list) else "N/A")
                            
                            st.markdown("---")
                            
                            suggestions_str = str(suggestions)
                            
                            st.markdown("### ✅ **Strengths**")
                            if "strength" in suggestions_str.lower():
                                strengths_section = [line.strip() for line in suggestions_str.split('\n') 
                                                   if any(keyword in line.lower() for keyword in ['strength', 'strong', 'good', 'excellent', 'well'])]
                                if strengths_section:
                                    for item in strengths_section[:5]:
                                        if item and len(item) > 5:
                                            st.markdown(f"• {item.lstrip('•-*0123456789.) ')}")
                                else:
                                    st.markdown("• Resume aligns well with job requirements")
                                    st.markdown("• Strong use of relevant keywords")
                            else:
                                st.markdown("• Resume aligns well with job requirements")
                                st.markdown("• Good keyword coverage")
                            
                            st.markdown("### ⚠️ **Weaknesses**")
                            if "weakness" in suggestions_str.lower() or "missing" in suggestions_str.lower():
                                weaknesses_section = [line.strip() for line in suggestions_str.split('\n') 
                                                    if any(keyword in line.lower() for keyword in ['weakness', 'weak', 'missing', 'lack', 'improve', 'add', 'consider'])]
                                if weaknesses_section:
                                    for item in weaknesses_section[:5]:
                                        if item and len(item) > 5:
                                            st.markdown(f"• {item.lstrip('•-*0123456789.) ')}")
                                else:
                                    st.markdown("• Consider adding more specific metrics")
                                    st.markdown("• Include more industry-specific keywords")
                            else:
                                st.markdown("• Consider adding more metrics and numbers")
                                st.markdown("• Include more technical keywords")
                            
                            st.markdown("### 🎯 **Missing Keywords**")
                            if isinstance(missing_keywords, list) and missing_keywords:
                                for keyword in missing_keywords[:10]:
                                    st.markdown(f"• {keyword}")
                            else:
                                st.markdown("• No critical keywords missing")
                            
                            st.markdown("### 💡 ATS Optimization Tips")
                            st.markdown("• Tailor the resume with keywords from the job description")
                            st.markdown("• Quantify achievements with numbers and metrics")
                            st.markdown("• Use strong action verbs to start bullet points")
                            st.markdown("• Use standard section headings (Work Experience, Education, Skills)")
                            st.markdown("• Avoid graphics, tables, and images")
                            st.markdown("• Avoid using information in headers and footers")
                            st.markdown("• Use standard fonts (Arial, Calibri, Times New Roman)")
                            st.markdown("• Save as .docx or .pdf for best compatibility")
                            st.markdown("• Proofread carefully to avoid errors")
                        
                        except:
                            st.info(response.content)
                    
                    st.markdown("---")

elif st.session_state.page == "AI Assistant":
    st.title("💬 AI Assistant")
    
    if not initial_api_key or not llm:
        st.error("⚠️ Please configure your Groq API key!")
    else:
        assist_mode = st.radio("Choose Mode:", ["General Chat", "Resume Guidance", "Interview Prep"], horizontal=True)
        
        if assist_mode == "General Chat":
            st.subheader("💬 Career Chat")
            
            for message in st.session_state.chat_history_general:
                if isinstance(message, dict) and "role" in message:
                    if message["role"] == "user":
                        st.markdown("""
                        <div style='background-color: #020e73; padding: 15px; border-radius: 8px; margin: 10px 0; border-left: 4px solid #1f77b4;'>
                            <h4><strong>👤 You:</strong></h4><br>
                            """ + message['content'] + """
                        </div>
                        """, unsafe_allow_html=True)
                    else:
                        st.markdown("""
                        <div style='background-color: #0f0f0f; padding: 15px; border-radius: 8px; margin: 10px 0; border-left: 4px solid #FFC107;'>
                            <h3><strong>🤖 Assistant:</strong></h3><br>
                            """ + message['content'] + """
                        </div>
                        """, unsafe_allow_html=True)
            
            col1, col2 = st.columns([0.85, 0.15])
            
            with col1:
                user_input = st.text_input("Ask anything about careers, resumes, or job hunting", key="ai_input_general_input")
            
            with col2:
                submit_button = st.button("Send", key="send_general")
            
            if submit_button and user_input:
                with st.spinner("Thinking..."):
                    st.session_state.chat_history_general.append({"role": "user", "content": user_input})
                    
                    chain = ChatPromptTemplate.from_messages([
                        ("system", "You are a helpful career coach. Answer questions about resumes, job hunting, and interviews."),
                        ("human", "{input}")
                    ]) | llm
                    
                    response = chain.invoke({"input": user_input})
                    st.session_state.chat_history_general.append({"role": "assistant", "content": response.content})
                    st.rerun()
            
            if st.session_state.chat_history_general and st.button("🗑️ Clear History", key="clear_general"):
                st.session_state.chat_history_general = []
                st.rerun()
        
        elif assist_mode == "Resume Guidance":
            st.subheader("📋 Resume-Specific Guidance")
            
            resume_choice = st.radio("Which resume?", ["Generated Resume", "Uploaded Resume"], horizontal=True)
            
            selected_resume = ""
            if resume_choice == "Generated Resume":
                selected_resume = st.session_state.generated_resume
            else:
                selected_resume = st.session_state.uploaded_resume
            
            if resume_choice == "Uploaded Resume" and not selected_resume:
                uploaded_file = st.file_uploader("Upload resume for guidance", type=['txt', 'pdf', 'docx'], key="guidance_uploader")
                if uploaded_file:
                    try:
                        if uploaded_file.type == "text/plain":
                            selected_resume = uploaded_file.read().decode("utf-8")
                        elif uploaded_file.type == "application/pdf" and PDF_SUPPORT:
                            pdf_reader = PyPDF2.PdfReader(uploaded_file)
                            selected_resume = "\n".join([page.extract_text() for page in pdf_reader.pages])
                        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" and DOCX_SUPPORT:
                            doc = DocxDocument(uploaded_file)
                            selected_resume = "\n".join([para.text for para in doc.paragraphs])
                        
                        st.session_state.uploaded_resume = selected_resume
                        st.success(f"✅ File loaded: {uploaded_file.name}")
                    except Exception as e:
                        st.error(f"Error reading file: {str(e)}")
            
            if not selected_resume:
                st.warning("⚠️ No resume found. Please generate a resume or upload one first.")
            else:
                st.success("✅ Resume loaded for guidance")
                
                for message in st.session_state.chat_history_guidance:
                    if isinstance(message, dict) and "role" in message:
                        if message["role"] == "user":
                            st.markdown("""
                            <div style='background-color: #04027a; padding: 15px; border-radius: 8px; margin: 10px 0; border-left: 4px solid #1f77b4;'>
                                <h4><strong>👤 You:</strong></h4><br>
                                """ + message['content'] + """
                            </div>
                            """, unsafe_allow_html=True)
                        else:
                            st.markdown("""
                            <div style='background-color: #0a0a0a; padding: 15px; border-radius: 8px; margin: 10px 0; border-left: 4px solid #FFC107;'>
                                <h4><strong>🤖 Assistant:</strong></h4><br>
                                """ + message['content'] + """
                            </div>
                            """, unsafe_allow_html=True)
                
                col1, col2 = st.columns([0.85, 0.15])
                
                with col1:
                    user_input = st.text_input("Ask about your resume", key="ai_input_guidance_input")
                
                with col2:
                    submit_button = st.button("Send", key="send_guidance")
                
                if submit_button and user_input:
                    with st.spinner("Analyzing..."):
                        st.session_state.chat_history_guidance.append({"role": "user", "content": user_input})
                        
                        chain = ChatPromptTemplate.from_messages([
                            ("system", "You are a resume expert. Provide specific, actionable feedback on resumes."),
                            ("human", "Resume:\n{resume}\n\nQuestion: {input}")
                        ]) | llm
                        
                        response = chain.invoke({"resume": selected_resume, "input": user_input})
                        st.session_state.chat_history_guidance.append({"role": "assistant", "content": response.content})
                        st.rerun()
                
                st.markdown("### 💡 Quick Actions")
                
                quick_questions = [
                    ("✨ Improve Writing", "Analyze my resume and suggest improvements for writing quality, clarity, and impact."),
                    ("🎯 Optimize Keywords", "Suggest important keywords and phrases to improve ATS compatibility."),
                    ("💪 Stronger Verbs", "Suggest stronger action verbs and more impactful phrasing for my bullet points."),
                    ("📊 Add Metrics", "Suggest where I can add metrics, numbers, and quantifiable achievements."),
                    ("❓ General Tips", "What are the most important things I should focus on in my resume?")
                ]
                
                cols = st.columns(5)
                for idx, (button_label, question_text) in enumerate(quick_questions):
                    with cols[idx]:
                        if st.button(button_label, use_container_width=True, key=f"quick_btn_{idx}"):
                            with st.spinner("Getting response..."):
                                st.session_state.chat_history_guidance.append({"role": "user", "content": question_text})
                                
                                chain = ChatPromptTemplate.from_messages([
                                    ("system", "You are a resume expert. Provide specific, actionable feedback."),
                                    ("human", "Resume:\n{resume}\n\nQuestion: {input}")
                                ]) | llm
                                
                                response = chain.invoke({"resume": selected_resume, "input": question_text})
                                st.session_state.chat_history_guidance.append({"role": "assistant", "content": response.content})
                                st.rerun()
                
                if st.session_state.chat_history_guidance and st.button("🗑️ Clear History", key="clear_guidance"):
                    st.session_state.chat_history_guidance = []
                    st.rerun()
        
        elif assist_mode == "Interview Prep":
            st.subheader("🎤 Interview Preparation")
            
            st.markdown("### 📚 Interview Resources")
            
            col1, col2 = st.columns(2)
            
            with col1:
                interview_type = st.selectbox(
                    "Select Interview Type:",
                    ["Behavioral", "Technical", "Situational", "Case Study", "Panel"]
                )
            
            with col2:
                industry = st.text_input("Your Industry/Role:", placeholder="e.g., Software Engineering, Marketing")
            
            col1, col2, col3 = st.columns(3)
            
            with col1:
                if st.button("❓ Generate Questions", use_container_width=True):
                    with st.spinner("Generating interview questions..."):
                        chain = ChatPromptTemplate.from_messages([
                            ("system", "You are an interview coach. Generate relevant interview questions."),
                            ("human", "Generate 5 {interview_type} interview questions for a {industry} role.")
                        ]) | llm
                        
                        response = chain.invoke({
                            "interview_type": interview_type,
                            "industry": industry or "general professional"
                        })
                        st.info(response.content)
            
            with col2:
                if st.button("💬 Practice STAR Method", use_container_width=True):
                    st.markdown("""
                    ### STAR Method Framework""")
                    
                    st.markdown("""**S - Situation**: Describe the situation/task""")
                    st.markdown("""**T - Task**: Explain your role or responsibility""")
                    st.markdown("""**A - Action**: Detail the specific actions you took""")
                    st.markdown("""**R - Result**: Share the measurable results achieved""")
                    
                    st.markdown("""**Example:** "In my role at [Company], I was tasked with [Task]. I [Action]. This resulted in [Result]."
                    """)
            
            with col3:
                if st.button("🎯 Mock Interview Tips", use_container_width=True):
                    st.success("""
                    ✅ **Interview Tips:**
                    - Research the company thoroughly
                    - Prepare 3-5 stories using STAR method
                    - Practice out loud before the interview
                    - Dress appropriately for the role
                    - Arrive/log in 5-10 minutes early
                    - Ask thoughtful questions about the role
                    - Send thank you email within 24 hours
                    - Mirror the interviewer's communication style
                    """)
            
            st.markdown("---")
            
            st.markdown("### 🎤 Practice Interview")
            
            practice_type = st.selectbox("Practice Question Type:", ["Random", "Behavioral", "Technical"], key="practice_type")
            
            for message in st.session_state.chat_history_interview:
                if isinstance(message, dict) and "role" in message:
                    if message["role"] == "user":
                        st.markdown("""
                        <div style='background-color: #04027a; padding: 15px; border-radius: 8px; margin: 10px 0; border-left: 4px solid #1f77b4;'>
                            <strong>👤 You:</strong><br>
                            """ + message['content'] + """
                        </div>
                        """, unsafe_allow_html=True)
                    else:
                        st.markdown("""
                        <div style='background-color: #0a0a0a; padding: 15px; border-radius: 8px; margin: 10px 0; border-left: 4px solid #4CAF50;'>
                            <strong>🎤 Interviewer:</strong><br>
                            """ + message['content'] + """
                        </div>
                        """, unsafe_allow_html=True)
            
            if st.button("🎤 Get Interview Question", use_container_width=True, key="get_question"):
                with st.spinner("Generating question..."):
                    chain = ChatPromptTemplate.from_messages([
                        ("system", f"You are a professional interviewer. Ask a {practice_type.lower()} interview question."),
                        ("human", "Ask me a {practice_type} interview question for {industry} role.")
                    ]) | llm
                    
                    response = chain.invoke({
                        "practice_type": practice_type,
                        "industry": industry or "general professional"
                    })
                    
                    st.session_state.chat_history_interview.append({
                        "role": "interviewer",
                        "content": response.content
                    })
                    st.rerun()
            
            if len(st.session_state.chat_history_interview) > 0 and st.session_state.chat_history_interview[-1].get("role") == "interviewer":
                user_answer = st.text_area("Your Answer:", height=150, key="interview_answer")
                
                if st.button("📝 Submit Answer & Get Feedback", use_container_width=True):
                    if user_answer.strip():
                        with st.spinner("Analyzing your answer..."):
                            st.session_state.chat_history_interview.append({
                                "role": "user",
                                "content": user_answer
                            })
                            
                            chain = ChatPromptTemplate.from_messages([
                                ("system", "You are an interview coach. Provide constructive feedback on interview answers."),
                                ("human", "Answer: {answer}\n\nProvide feedback on: clarity, relevance, use of STAR method, confidence indicators.")
                            ]) | llm
                            
                            response = chain.invoke({"answer": user_answer})
                            
                            st.session_state.chat_history_interview.append({
                                "role": "interviewer",
                                "content": f"**Feedback:** {response.content}"
                            })
                            st.rerun()
                    else:
                        st.error("Please provide an answer before submitting.")
            
            if st.session_state.chat_history_interview and st.button("🗑️ Clear Practice Session", key="clear_interview"):
                st.session_state.chat_history_interview = []
                st.rerun()
